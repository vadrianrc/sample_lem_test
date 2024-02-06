# -*- coding: utf-8 -*-

from flask import Flask
app = Flask(__name__)



@app.route('/execute-python', methods=['GET'])
def execute_python():
    # Your Python code logic goes here
    from google.oauth2.credentials import Credentials
    from googleapiclient.discovery import build
    from google.oauth2 import service_account
    import pandas as pd
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from datetime import datetime
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn
    import re

    SCOPES = ['https://www.googleapis.com/auth/spreadsheets']
    KEY = 'key.json'
    SPREADSHEET_ID = '16M4fktVRMrbowkq4LHP_R6t66VGn-6qZMWkbsW_nu2Q'

    creds = None
    creds = service_account.Credentials.from_service_account_file(KEY, scopes=SCOPES)

    service = build('sheets', 'v4', credentials=creds)
    sheet = service.spreadsheets()

    result = sheet.values().get(spreadsheetId=SPREADSHEET_ID, range='Hoja 1!B7:L').execute()
    values = result.get('values',[])
    # print(values)
    df = pd.DataFrame(values)
    df.columns = df.iloc[0]
    df = df[1:]



    ###################################################################################

    def add_signature_lines(footer, signature_left=None, signature_right=None):
        footer.add_paragraph("V°B° __________________________")
        footer.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT

        run = footer.paragraphs[-1].add_run("\t\t\t________________________")
        run.font.size = Pt(10)
        run.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if signature_left:
            footer.add_paragraph(signature_left).alignment = WD_ALIGN_PARAGRAPH.LEFT

        if signature_right:
            run = footer.paragraphs[-1].add_run(signature_right)
            run.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def add_footer(doc, footer_text):
        section = doc.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0]
        run = paragraph.add_run(footer_text)
        run.font.size = Pt(8)


    def insertHR(paragraph):
        p = paragraph._p  # p is the <w:p> XML element
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pPr.insert_element_before(pBdr,
            'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
            'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
            'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
            'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
            'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
            'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
            'w:pPrChange'
        )
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)

    def strWriteDate():
        meses = {1:'Enero', 2:'Febrero', 3:'Marzo', 4:'Abril', 5:'Mayo', 6:'Junio', 7:'Julio', 8:'Agosto', 9:'Septiembre', 10:'Octubre', 11:'Noviembre', 12:'Diciembre'}
        date = datetime.now()
        d_month = meses.get(int(date.strftime("%m")))
        d_year = date.strftime("%Y")
        strDate =f"Lima, {d_month} del {d_year}"
        return strDate


    BOSS_NAME = 'Ing. Oscar Miranda Hospinal'
    BOSS_DESCRP = 'Jefe del Laboratorio N°1 Ensayo de Materiales'

    SUP_NAME = 'Lic. Ladislao Jesús Basurto Pinao'
    SUP_DESCRP = 'Asistente del Laboratorio N°1 Ensayo de Materiales'

    CCEC = 'Realización de Supervisión de ensayos de Laboratorio, correspondiente'
    CCEC_DATE = "al mes de MONTH del YEAR"

    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    # Title
    title = document.add_paragraph()
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style = document.styles['Normal']
    run = title.add_run('INFORME')
    run.bold = True
    run.italic = True
    run.underline = True

    # Requests
    requests = document.add_paragraph()
    requests.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = requests.add_run('A\t\t')
    run = requests.add_run(':\t'+ BOSS_NAME +"\n")
    run = requests.add_run('\t\t\t' + BOSS_DESCRP)
    requests = document.add_paragraph()
    run = requests.add_run('De\t\t')
    run = requests.add_run(':\t'+ SUP_NAME +"\n")
    run = requests.add_run('\t\t\t' + SUP_DESCRP)
    requests = document.add_paragraph()
    run = requests.add_run('Asunto\t\t')
    run = requests.add_run(':\t'+ CCEC +"\n")
    run = requests.add_run('\t\t\t' + CCEC_DATE)
    requests = document.add_paragraph()
    run = requests.add_run('Fecha\t\t')
    run = requests.add_run(':\t'+ strWriteDate() + "\n")

    df_datosfinales = df.groupby('Cliente')['Descripción del Ensayo'].agg(list).reset_index()

    print(df_datosfinales)

    insertHR(requests)
    body = document.add_paragraph()
    run = body.add_run('\nActividades desarrolladas:\n')

    for ind in df_datosfinales.index:
        # document.add_paragraph(df_datosfinales['Cliente'][ind], style='List Number').bold = True
        body_cliente = document.add_paragraph(style='List Number')
        body_cliente.add_run(df_datosfinales['Cliente'][ind]).bold = True

        body_ensayos = document.add_paragraph()
        for ensayo in df_datosfinales['Descripción del Ensayo'][ind]:
            body_ensayos.add_run(ensayo+"\n")
            # document.add_paragraph().add_run(ensayo)


    add_footer(document, None)

    for section in document.sections:
            footer = section.footer
            add_signature_lines(footer, "\tIng. Oscar Miranda Hospinal", "\t\tLic. LADISLAO JESÚS BASURTO PINAO")



    # clientes_unicos = df.drop_duplicates(subset='Cliente')
    # print(clientes_unicos)

    # for _, cliente in clientes_unicos.iterrows():
    #     # Filtra el DataFrame original para el cliente actual
    #     cliente_actual = df[df['Cliente'] == cliente['Cliente']]
        
    #     # Ordena por fecha
    #     cliente_actual = cliente_actual.sort_values(by='Fec. de Ingreso')

    #     # Imprime la información del cliente y sus pedidos ordenados por fecha
    #     document.add_paragraph(cliente['Cliente'], style='List Number') 
    #     datos = document.add_paragraph()
    #     # datos.add_run(cliente['Fec. de Ingreso'] + "\n").bold = True
    #     # datos.add_run(patron.sub('',  cliente['Descripción del Ensayo']))

    #     datos.add_run(cliente['Descripción del Ensayo'])


    document.save('demo.docx')
    ###################################################################################




    result = "Python code executed successfully!"
    return result

if __name__ == '__main__':
    app.run(debug=True)
